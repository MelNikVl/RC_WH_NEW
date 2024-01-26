import datetime
import os
from typing import Dict
import fastapi
import xlsxwriter
from fastapi import Depends, Request, Response
from sqlalchemy import desc
from sqlalchemy.orm import Session
from fastapi.templating import Jinja2Templates
from fastapi.encoders import jsonable_encoder
from crud.geolocation import GeoLocationCRUD
from crud.materials import MaterialCRUD
from starlette import status
from app.controllers.materials import user_dependency
from app.utils.auth import AuthUtil
from app.utils.utils import get_first_photo, response, Mail
from db.db import get_db
from models.models import User, GeoLocation, Material, Repair, Notifications, LogItem, Raw_1c, Email
from app.payload.request import InvoiceCreateRequest, MaterialsListRequest
from docx import Document
from fastapi.responses import FileResponse

templates = Jinja2Templates(directory="templates")


class FrontMainController:

    # эндпоинт что бы отслеживать валидность токена
    @staticmethod
    async def ping(user: user_dependency):
        return Response("ok", media_type="text/plain")

    @staticmethod
    async def generate_list(materials: MaterialsListRequest) -> FileResponse:
        out_filename = 'list.xlsx'
        headers = ["ID", "Category", "Title", "Description", "Creation date", "Place", "Status"]
        workbook = xlsxwriter.Workbook(out_filename)
        worksheet = workbook.add_worksheet()
        cell_format = workbook.add_format({"valign": "top", 'text_wrap': True})
        for col_num, data in enumerate(headers):
            worksheet.write(0, col_num, data)
        counter = 1

        for item in materials.data:
            for col_num, data in enumerate([item[0], item[1], item[2], item[3], item[4], item[5], item[6]]):
                worksheet.write(counter, col_num, data, cell_format)
                worksheet.set_column(counter, col_num, 30)
            counter += 1

        workbook.close()

        # with open(out_filename, 'w', newline='') as csvfile:
        #     out_file = csv.writer(csvfile)
        #     out_file.writerow(["ID", "Category", "Title", "Description", "Creation date", "Place", "Status"])
        #     for item in materials.data:
        #         out_file.writerow([item[0], item[1], item[2], item[3], item[4], item[5], item[6]], )
        return FileResponse(out_filename)

    @staticmethod
    async def generate_invoice(
            materials: InvoiceCreateRequest) -> FileResponse:  # генерируем ворд файл из таблицы списания
        directory = "invoices"
        if not os.path.exists(directory):
            os.makedirs(directory)
        out_name = os.path.join(directory, datetime.datetime.now().strftime("%Y-%m-%d %H-%M-%S") + ".docx")

        document = Document()
        par = document.add_paragraph()
        par.add_run('Акт на списание Материальных средств ЗАО «РЕНЕЙССАНС КОНСТРАКШН»').bold = True
        par.alignment = 1

        par = document.add_paragraph()
        par.add_run('(компьютеры, оргтеника, периферийное оборудование)').bold = True
        par.alignment = 1

        par = document.add_paragraph()
        par.add_run('№_____________ от ' + datetime.datetime.now().strftime("%Y.%m.%d")).bold = True
        par.alignment = 1

        table = document.add_table(1, 6)
        table.style = 'Table Grid'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = '№'
        heading_cells[1].text = 'ID'
        heading_cells[2].text = 'Category'
        heading_cells[3].text = 'Title'
        heading_cells[4].text = 'Description'
        heading_cells[5].text = 'Date'

        for item in materials.data:
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]
            row_cells[3].text = item[3]
            row_cells[4].text = item[4]
            row_cells[5].text = item[5]

        par = document.add_paragraph()
        par.add_run('Подпись руководителя IT отдела ______________________________').bold = True
        par.alignment = 1

        document.save(out_name)
        return FileResponse(out_name)

    @staticmethod
    async def index(db: Session = Depends(get_db),
                    request: Request = None,
                    t: str = None):

        out: Dict = {}
        materials = await MaterialCRUD.list_of_materials(db=db)

        # добавляем фото
        for i in materials: i.pic = get_first_photo(i.id)["picture"]

        out[0] = jsonable_encoder(materials)

        # подсчет активов
        out["count_warehouse"] = await GeoLocationCRUD.get_materials_at_warehouses(db=db)

        try:
            result = await AuthUtil.decode_jwt(t)
            out["username"] = result["username"]
            out["role"] = result["role"]
        except Exception as e:
            print(e)
            return fastapi.responses.RedirectResponse('/app/auth', status_code=status.HTTP_301_MOVED_PERMANENTLY)

        # добавляем местоположения
        for i in range(len(out[0])):
            geolocation_place = await GeoLocationCRUD.current_place(material_id=out[0][i]['id'], db=db)
            if geolocation_place is None:
                out[0][i].update({"geolocation_place": "место не указано"})
                out[0][i].update({"geolocation_status": "статуса нет"})
            else:
                out[0][i].update({"geolocation_place": geolocation_place.place})
                out[0][i].update({"geolocation_status": geolocation_place.status})

        # добавляем почты для уведомлений
        emails_1 = []
        for i in Mail.get_emails(db):
            emails_1.append(i.addr)

        fik = jsonable_encoder(emails_1)
        out["emails_to_nothification_main_table"] = fik

        out["token"] = t

        return templates.TemplateResponse("table.html", {"request": request, "data": out})

    @staticmethod
    async def user_auth(db: Session = Depends(get_db),
                        request: Request = None,
                        ):
        if len(db.query(User).all()) == 0:
            user = User(chat_id=11111,
                        username="admin",
                        password="$2b$12$LHizNG913MQ.FlcjXS9eGufsJK2yp5xdbt6dvCGzOnMQbCSrLna2.",
                        is_admin=True)
            db.add(user)
            db.commit()
        return templates.TemplateResponse("auth.html", {"request": request})

    @staticmethod
    async def admins_page(db: Session = Depends(get_db),
                          request: Request = None,
                          t: str = None  # jwt токен
                          ):
        is_admin = 0
        try:
            result = await AuthUtil.decode_jwt(t)
            is_admin = int(result["role"])
        except Exception as e:
            return fastapi.responses.RedirectResponse('/app/auth', status_code=status.HTTP_301_MOVED_PERMANENTLY)

        if is_admin:
            out: Dict = {"token": t,
                         "users": jsonable_encoder(db.query(User).all()),
                         "count": db.query(User).count(),
                         "admins_count": db.query(User).filter(User.is_admin == True).count()}
            result = await AuthUtil.decode_jwt(t)

            out["username"] = result["username"]
            out["role"] = result["role"]
            out["email_for_noth"] = Mail.get_emails(db)

            return templates.TemplateResponse("admins_page.html", {"request": request, "data": out})
        else:
            return fastapi.responses.RedirectResponse('/app', status_code=status.HTTP_301_MOVED_PERMANENTLY)

    @staticmethod
    async def instructions(db: Session = Depends(get_db),
                           request: Request = None,
                           t: str = None  # jwt токен
                           ):
        try:
            result = await AuthUtil.decode_jwt(t)
        except Exception as e:
            return fastapi.responses.RedirectResponse('/app/auth', status_code=status.HTTP_301_MOVED_PERMANENTLY)

        out: Dict = {"token": t,
                     "username": result["username"],
                     "role": result["role"]}

        return templates.TemplateResponse("instructions.html", {"request": request, "data": out})

    @staticmethod
    async def only_one_card(material_id,
                            request: Request = None,
                            t: str = None,  # jwt токен
                            db: Session = Depends(get_db)):
        try:
            result = await AuthUtil.decode_jwt(t)
        except Exception as e:
            print(e)
            return fastapi.responses.RedirectResponse('/app/auth', status_code=status.HTTP_301_MOVED_PERMANENTLY)

        material_card = jsonable_encoder(db.query(Material).filter(Material.id == material_id).first())
        material_geo1 = db.query(GeoLocation).filter(GeoLocation.material_id == material_id).filter(GeoLocation.geo_type != 2).order_by(
            desc(GeoLocation.date_time)).all()

        for i in material_geo1:
            current_geo = db.query(GeoLocation).filter(GeoLocation.material_id == material_id).order_by(desc(GeoLocation.date_time)).all()[0]

        for item in material_geo1:
            item.date_time = item.date_time.strftime("%Y-%m-%d")

        datetime_obj = datetime.datetime.strptime(material_card['date_time'], "%Y-%m-%dT%H:%M:%S.%f")
        formatted_date_time = datetime_obj.strftime("%Y-%m-%d %H:%M")
        raw_1c = jsonable_encoder(db.query(Raw_1c).filter(Raw_1c.material_id == material_id).first())
        formatted_time_1c = db.query(Raw_1c).filter(Raw_1c.material_id == material_id).first().date_time.strftime("%d %B %Y")

        emails_to_nothification_one_card = db.query(Email).all()
        emails_1 = []
        for i in emails_to_nothification_one_card:
            emails_1.append(i.addr)

        fik = jsonable_encoder(emails_1)

        out: Dict = {"username": result["username"],
                     "token": t, "one_material": material_card,
                     "geo_material": material_geo1,
                     "repairs": GeoLocationCRUD.list_of_repair(material_id, db),
                     "date_time_f": formatted_date_time,
                     "current_place": current_geo.place,
                     "current_user": current_geo.client_mail,
                     "current_status": current_geo.status,
                     "photo": get_first_photo(material_id)["picture"],
                     "len_of_files": get_first_photo(material_id)["len_of_files"],
                     "material_id": str(material_id),
                     "role": result["role"],
                     "comments": await MaterialCRUD.get_comments(material_id, db),
                     "raw_1c": raw_1c,
                     "formatted_time_1c": formatted_time_1c,
                     "emails_to_nothification_one_card": fik}

        return templates.TemplateResponse("one_material.html", {"request": request, "data": out})

    @staticmethod
    async def repairs_page(request: Request = None,
                           t: str = None,  # jwt токен
                           db: Session = Depends(get_db)):

        try:
            result = await AuthUtil.decode_jwt(t)
        except Exception as e:
            return fastapi.responses.RedirectResponse('/app/auth', status_code=status.HTTP_301_MOVED_PERMANENTLY)

        products = db.query(Repair).filter(Repair.id.in_(GeoLocationCRUD.list_of_active_repair(db))).all()

        out: Dict = {"token": t,
                     "actives_in_repair": products,
                     "count_repair": len(products),
                     "role": result["role"],
                     "username": result["username"]}

        return templates.TemplateResponse("repair_page.html", {"request": request, "data": out})

    @staticmethod
    async def trash_page(db: Session = Depends(get_db),
                         request: Request = None,
                         t: str = None  # jwt токен
                         ):
        # проверка токена на валидность и если он не вализный - переадресация на авторизацию
        try:
            result = await AuthUtil.decode_jwt(t)
        except Exception as e:
            return fastapi.responses.RedirectResponse('/app/auth', status_code=status.HTTP_301_MOVED_PERMANENTLY)
        out: Dict = {}
        materials_for_trash = await GeoLocationCRUD.get_materials_for_trash(db=db)

        for item in materials_for_trash:
            item.date_time = item.date_time.strftime("%Y-%m-%d")

        emails_to_nothification_one_card = db.query(Email).all()
        emails_1 = []
        for i in emails_to_nothification_one_card: emails_1.append(i.addr)
        fik = jsonable_encoder(emails_1)

        out[0] = materials_for_trash
        out["token"] = t
        out["count_for_trash"] = len(materials_for_trash)
        out["role"] = result["role"]
        out["username"] = result["username"]
        out["notifications"] = fik

        return templates.TemplateResponse("trash_page.html", {"request": request, "data": out})

    @staticmethod
    async def notification_answer(unique_code,
                                  material_id,
                                  db: Session = Depends(get_db),
                                  request: Request = None
                                  ):
        answer = db.query(Notifications).filter(Notifications.unique_code == unique_code).first()
        answer.read = True
        db.commit()

        create_event = LogItem(kind_table="Уведомления",
                               user_id=0,
                               passive_id=material_id,
                               modified_cols="ответ на уведомление",
                               values_of_change=f'пользователь {answer.user} '
                                                f'подтвердил перемещение актива - {material_id}, '
                                                f' идентификатор уведомления - {answer.unique_code}',
                               date_time=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                               )
        db.add(create_event)
        db.commit()

        return response(data="Спасибо за ответ. Уведомление получено. Вы можете закрыть эту страницу", status=True)

    @staticmethod
    async def notifications_page(db: Session = Depends(get_db),
                                 request: Request = None,
                                 t: str = None  # jwt токен,
                                 ):
        # проверка токена на валидность и если он не вализный - переадресация на авторизацию
        try:
            result = await AuthUtil.decode_jwt(t)
        except Exception as e:
            return fastapi.responses.RedirectResponse('/app/auth', status_code=status.HTTP_301_MOVED_PERMANENTLY)

        out: Dict = {}
        notifications_all = db.query(Notifications).order_by(Notifications.date_time.desc()).all()

        for item in notifications_all:
            item.date_time = item.date_time.strftime("%Y-%m-%d")

        out["token"] = t
        out["notifications"] = notifications_all

        return templates.TemplateResponse("notifications.html", {"request": request, "data": out})

    @staticmethod
    async def test_page(db: Session = Depends(get_db),
                        request: Request = None,
                        t: str = None  # jwt токен,
                        ):
        # проверка токена на валидность и если он не вализный - переадресация на авторизацию
        try:
            result = await AuthUtil.decode_jwt(t)
        except Exception as e:
            return fastapi.responses.RedirectResponse('/app/auth', status_code=status.HTTP_301_MOVED_PERMANENTLY)

        out: Dict = {"token": t}

        return templates.TemplateResponse("test.html", {"request": request, "data": out})


    @staticmethod
    async def only_1c(db: Session = Depends(get_db),
                      request: Request = None,
                      t: str = None  # jwt токен,
                      ):
        try:
            result = await AuthUtil.decode_jwt(t)
        except Exception as e:
            return fastapi.responses.RedirectResponse('/app/auth', status_code=status.HTTP_301_MOVED_PERMANENTLY)

        out: Dict = {"token": t}

        return templates.TemplateResponse("only_1c.html", {"request": request, "data": out})
